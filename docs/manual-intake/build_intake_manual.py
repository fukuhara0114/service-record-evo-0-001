# -*- coding: utf-8 -*-
"""Build Intake操作マニュアル.docx from screenshots."""
from __future__ import annotations

import os
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

SCRIPT_DIR = Path(__file__).resolve().parent
PROJECT_DIR = SCRIPT_DIR.parent.parent
IMG_DIR = SCRIPT_DIR
OUT_PATH = PROJECT_DIR / "docs" / "Intake操作マニュアル.docx"
BASE_URL = "http://localhost/dev/git_test/service-record-evo-0-001"

JP_FONTS = ("游ゴシック", "Yu Gothic", "メイリオ", "Meiryo", "ＭＳ ゴシック", "MS Gothic")


def pick_jp_font() -> str | None:
    windir = Path(os.environ.get("WINDIR", r"C:\Windows"))
    fonts = windir / "Fonts"
    candidates = [
        ("Yu Gothic", ["YuGothM.ttc", "YuGothR.ttc", "yugothm.ttc"]),
        ("Meiryo", ["meiryo.ttc", "Meiryo.ttc"]),
        ("MS Gothic", ["msgothic.ttc", "MSGOTHIC.TTC"]),
    ]
    for name, files in candidates:
        for f in files:
            if (fonts / f).exists():
                return name
    return None


JP_FONT = pick_jp_font()
FIGURE_NO = 0
EMBEDDED: list[str] = []
MISSING: list[str] = []


def set_run_font(run, size_pt: float | None = None, bold: bool | None = None):
    if bold is not None:
        run.bold = bold
    if size_pt is not None:
        run.font.size = Pt(size_pt)
    if JP_FONT:
        run.font.name = JP_FONT
        r = run._element
        rPr = r.get_or_add_rPr()
        rFonts = rPr.get_or_add_rFonts()
        rFonts.set(qn("w:eastAsia"), JP_FONT)


def set_paragraph_font(paragraph, size_pt: float = 11):
    for run in paragraph.runs:
        set_run_font(run, size_pt=size_pt)


def add_heading_jp(doc: Document, text: str, level: int = 1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        set_run_font(run)
    return h


def add_para(doc: Document, text: str, *, bold: bool = False, size: float = 11):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_run_font(run, size_pt=size, bold=bold)
    return p


def add_bullets(doc: Document, items: list[str]):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(item)
        set_run_font(run, size_pt=11)


def add_numbered(doc: Document, items: list[str]):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        run = p.add_run(item)
        set_run_font(run, size_pt=11)


def add_figure(doc: Document, filename: str, caption: str):
    global FIGURE_NO
    path = IMG_DIR / filename
    FIGURE_NO += 1
    cap_text = f"図{FIGURE_NO}: {caption}"
    if not path.exists():
        MISSING.append(filename)
        add_para(
            doc,
            f"（画像なし: {filename} が見つかりませんでした）",
            bold=True,
            size=10,
        )
        add_para(doc, cap_text, size=10)
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(path), width=Inches(6.0))
    EMBEDDED.append(filename)
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cap.add_run(cap_text)
    set_run_font(r, size_pt=10)


def add_table(doc: Document, headers: list[str], rows: list[list[str]]):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ""
        p = hdr[i].paragraphs[0]
        run = p.add_run(h)
        set_run_font(run, size_pt=10, bold=True)
    for ri, row in enumerate(rows):
        cells = table.rows[ri + 1].cells
        for ci, val in enumerate(row):
            cells[ci].text = ""
            p = cells[ci].paragraphs[0]
            run = p.add_run(val)
            set_run_font(run, size_pt=10)
    doc.add_paragraph()


def configure_styles(doc: Document):
    style = doc.styles["Normal"]
    style.font.size = Pt(11)
    if JP_FONT:
        style.font.name = JP_FONT
        style.element.rPr.rFonts.set(qn("w:eastAsia"), JP_FONT)
    for name in ("Heading 1", "Heading 2", "Heading 3"):
        if name in doc.styles:
            hs = doc.styles[name]
            if JP_FONT:
                hs.font.name = JP_FONT
                try:
                    hs.element.rPr.rFonts.set(qn("w:eastAsia"), JP_FONT)
                except Exception:
                    pass


def build():
    doc = Document()
    configure_styles(doc)

    # Cover
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("Intake操作マニュアル")
    set_run_font(r, size_pt=28, bold=True)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run("管理者向け（案件登録／Intake）")
    set_run_font(r, size_pt=16)

    meta_lines = [
        "日付: 2026-08-21",
        "対象読者: 管理者（administrator）",
        f"ベースURL: {BASE_URL}",
    ]
    for line in meta_lines:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(line)
        set_run_font(r, size_pt=12)

    doc.add_paragraph()

    # 1. はじめに
    add_heading_jp(doc, "1. はじめに", 1)

    add_heading_jp(doc, "1.1 対象読者", 2)
    add_para(
        doc,
        "本マニュアルは、サービスレコード管理システムにおける Intake（案件登録）機能を利用する管理者（administrator）向けです。"
        "案件の新規作成、未登録ファイルからの登録、既存案件へのPDF紐づけ、Loaner案件との関連付けなど、日常運用で必要な操作を説明します。",
    )

    add_heading_jp(doc, "1.2 前提条件", 2)
    add_bullets(
        doc,
        [
            "管理者アカウントでログイン済みであること。",
            f"ブラウザから {BASE_URL} にアクセスできること。",
            "案件登録に必要なマスタ（機種・ディーラー等）が整備されていること。",
            "ファイルから登録する場合は、PDF等の申請書類が手元にあること。",
        ],
    )

    add_heading_jp(doc, "1.3 用語", 2)
    add_bullets(
        doc,
        [
            "Intake: 未登録ファイルの取り込みと案件新規作成を行う機能画面群。",
            "サービス案件: 通常のサービス対応案件（service）。",
            "Loaner案件: 貸出機（loaner）に関する案件。",
            "未登録ファイル: アップロード済みだが、まだ案件に紐づいていない書類。",
            "OCR: 申請フォームPDF等から文字情報を読み取る機能。",
            "parentID: service案件が参照する親となるloaner案件の識別子。",
        ],
    )

    # 2. 画面への入り方
    add_heading_jp(doc, "2. 画面への入り方（ホーム → Intake）", 1)
    add_para(
        doc,
        "ログイン後、ホーム画面から Intake へ進みます。ホーム上の Intake への導線（メニュー／カード等）を選択してください。",
    )
    add_numbered(
        doc,
        [
            "ブラウザでベースURLにアクセスし、管理者としてログインする。",
            "ホーム画面を表示する。",
            "Intake（案件登録）へのリンクまたはカードを選択する。",
            "未登録ファイル一覧（Intake一覧）が表示されることを確認する。",
        ],
    )
    add_figure(doc, "01-home.png", "ホーム画面")

    # 3. 未登録ファイル一覧
    add_heading_jp(doc, "3. 未登録ファイル一覧の見方", 1)
    add_para(
        doc,
        "Intakeの入口画面では、未登録ファイルの一覧と、案件作成のための3つの入口カードが表示されます。"
        "一覧から既存の未登録ファイルを選んで案件化したり、ファイルなし／ファイル追加で新規作成を開始できます。",
    )
    add_figure(doc, "02-intake-list.png", "未登録ファイル一覧（Intake）")

    add_heading_jp(doc, "3.1 3入口カードの説明", 2)
    add_bullets(
        doc,
        [
            "添付なしで新規登録: ファイルを付けずに、サービス／Loanerの空の案件作成画面を開く。",
            "ファイル追加（選択／ドラッグ＆ドロップ）: PC上のファイルをアップロードし、そのファイル付きで新規案件を作成する。",
            "既存未登録ファイルから「このファイルで新規登録」: 一覧上の未登録ファイルを使い、申請フォーム付きの新規案件を作成する（OCR利用可）。",
        ],
    )
    add_para(
        doc,
        "作成画面では、さらに「既存案件（サービス案件）検索」や「loaner案件検索」により、既存案件への紐づけや親Loanerの指定も行えます（詳細は第5章・第6章）。",
    )

    # 4. 案件種類ダイアログ
    add_heading_jp(doc, "4. 案件種類ダイアログ（サービス／Loaner）", 1)
    add_para(
        doc,
        "新規作成を開始すると、案件種類を選択するダイアログが表示されます。"
        "「サービス」または「Loaner」を選び、以降の作成画面の種類が決まります。"
        "添付なしの場合と、ファイル追加（アップロード）の場合の双方で、同様の種類選択が行われます。",
    )
    add_figure(doc, "03-case-type-dialog.png", "案件種類ダイアログ（添付なし）")
    add_figure(doc, "11-case-type-upload.png", "案件種類ダイアログ（ファイル追加時）")
    add_bullets(
        doc,
        [
            "サービス: サービス案件の作成フローへ進む。",
            "Loaner: Loaner案件の作成フローへ進む。",
            "キャンセルすると、未登録ファイル一覧に戻る（またはダイアログを閉じる）。",
        ],
    )

    # 5. 作成パターン一覧表
    add_heading_jp(doc, "5. 作成パターン一覧表（マトリクス）", 1)
    add_para(
        doc,
        "Intakeからの案件作成・紐づけには、入口と案件種類の組み合わせで次のパターンがあります。"
        "運用時は、目的に応じて該当パターンを選んでください。",
    )
    add_table(
        doc,
        ["#", "入口", "案件種類", "結果"],
        [
            ["1", "添付なし", "サービス", "新規 service 案件"],
            ["2", "添付なし", "Loaner", "新規 loaner 案件"],
            ["3", "ファイル追加(選択/D&D)", "サービス", "アップロード後、ファイル付き新規 service"],
            ["4", "ファイル追加", "Loaner", "アップロード後、ファイル付き新規 loaner"],
            [
                "5",
                "既存未登録ファイル「このファイルで新規登録」",
                "サービス",
                "申請フォーム付き新規 service（OCR可）",
            ],
            ["6", "同上", "Loaner", "申請フォーム付き新規 loaner"],
            [
                "7",
                "作成画面「既存案件(サービス案件)検索」",
                "—",
                "選択案件へ PDF アタッチ（ソースファイル必須。添付なしでは不可）",
            ],
            [
                "8",
                "作成画面「loaner案件検索」",
                "—",
                "新規 service 作成し選択 loaner の parentID に設定",
            ],
        ],
    )

    # 6. パターン詳細
    add_heading_jp(doc, "6. パターン詳細", 1)

    add_heading_jp(doc, "6.1 パターン1: 添付なしでサービス案件を新規作成", 2)
    add_numbered(
        doc,
        [
            "未登録ファイル一覧で「添付なしで新規登録」を選ぶ。",
            "案件種類ダイアログで「サービス」を選択する。",
            "サービス案件の作成画面が開く（空のフォーム）。",
            "必要項目（顧客情報、機種、ディーラー、内容など）を入力する。",
            "機種マスタ選択など共通操作で値を確定する（第7章参照）。",
            "保存する。保存後は未登録ファイル一覧へ戻る。",
        ],
    )
    add_figure(doc, "04-create-service-blank.png", "添付なしサービス案件の作成画面")
    add_figure(doc, "10-master-select.png", "マスタ選択（機種など）")

    add_heading_jp(doc, "6.2 パターン2: 添付なしでLoaner案件を新規作成", 2)
    add_numbered(
        doc,
        [
            "未登録ファイル一覧で「添付なしで新規登録」を選ぶ。",
            "案件種類ダイアログで「Loaner」を選択する。",
            "Loaner案件の作成画面が開く。",
            "貸出機に関する必要項目を入力する。",
            "保存する。保存後は未登録ファイル一覧へ戻る。",
        ],
    )
    add_figure(doc, "08-create-loaner-blank.png", "Loaner新規作成画面")

    add_heading_jp(doc, "6.3 パターン3・4: ファイル追加から新規作成", 2)
    add_para(
        doc,
        "ファイル選択またはドラッグ＆ドロップでファイルを追加し、案件種類（サービス／Loaner）を選ぶと、"
        "アップロード後にそのファイルが付いた新規案件の作成画面へ進みます。手順の流れはパターン1・2と同様で、冒頭にアップロードが加わります。",
    )
    add_numbered(
        doc,
        [
            "「ファイル追加」カードでファイルを選択するか、D&Dする。",
            "案件種類ダイアログで「サービス」または「Loaner」を選ぶ。",
            "アップロード完了後、ファイル付きの作成画面で項目を入力し保存する。",
        ],
    )

    add_heading_jp(doc, "6.4 パターン5・6: 既存未登録ファイルから新規登録（OCR可）", 2)
    add_para(
        doc,
        "一覧の未登録ファイルに対し「このファイルで新規登録」を実行すると、申請フォーム付きの新規案件を作成できます。"
        "サービス／Loanerいずれも選択可能です。申請フォームがある場合はOCR読取を利用できます。",
    )
    add_numbered(
        doc,
        [
            "未登録ファイル一覧から対象ファイルを確認する。",
            "「このファイルで新規登録」を実行する。",
            "案件種類（サービス／Loaner）を選択する。",
            "作成画面で申請フォームを確認し、必要に応じてOCR読取を実行する。",
            "OCR結果を確認・修正し、拡大・回転などでプレビューを見やすくする。",
            "関連する他の未登録書類があれば確認し、必要に応じて取り込む。",
            "項目を確定して保存する。",
        ],
    )
    add_figure(doc, "09-create-from-file-service.png", "ファイルからサービス案件作成（OCR・拡大回転）")
    add_figure(doc, "07-related-files.png", "関連未登録書類")

    add_heading_jp(doc, "6.5 パターン7: 既存サービス案件へPDFをアタッチ", 2)
    add_para(
        doc,
        "作成画面の「既存案件（サービス案件）検索」から既存のサービス案件を選び、ソースとなるファイルをその案件へ添付できます。"
        "ソースファイルが必須のため、添付なしで開始したフローではこの操作はできません。",
    )
    add_numbered(
        doc,
        [
            "ファイル付きでIntakeの作成フローを開始する（ファイル追加または既存未登録ファイルから）。",
            "作成画面で「既存案件（サービス案件）検索」を開く。",
            "対象のサービス案件を検索・選択する。",
            "ソースファイルが当該案件へPDFアタッチされることを確認する。",
            "必要に応じて内容を確認し、処理を完了する。",
        ],
    )
    add_figure(doc, "05-existing-service-search.png", "既存サービス案件検索・紐づけ")

    add_heading_jp(doc, "6.6 パターン8: loaner案件へ紐づけ（parentID設定）", 2)
    add_para(
        doc,
        "作成画面の「loaner案件検索」で既存のLoaner案件を選択すると、新規のサービス案件を作成し、"
        "選択したLoanerの parentID に設定する形で関連付けます。",
    )
    add_numbered(
        doc,
        [
            "サービス案件の作成画面を開く（必要に応じてファイル付き）。",
            "「loaner案件検索」を開く。",
            "紐づけたいLoaner案件を検索・選択する。",
            "新規 service が作成され、選択Loanerが parentID として設定されることを確認する。",
            "残りの項目を入力して保存する。",
        ],
    )
    add_figure(doc, "06-loaner-link-search.png", "loaner案件検索・紐づけ")

    # 7. 共通操作
    add_heading_jp(doc, "7. 作成画面の共通操作", 1)
    add_para(doc, "サービス／Loanerいずれの作成画面でも、次の共通操作を把握しておくとスムーズです。")
    add_bullets(
        doc,
        [
            "タブ: 入力項目がタブ分けされている場合は、必要なタブを切り替えて漏れなく入力する。",
            "機種選択: マスタから機種を検索・選択する（図「マスタ選択」参照）。",
            "dealer選択: ディーラー（販売店等）をマスタから選択する。",
            "switch: 画面上のスイッチ（トグル）で表示・入力モードや関連オプションを切り替える。",
            "保守検索: 保守情報や関連レコードを検索し、案件に反映する。",
            "保存後の遷移: 保存に成功すると、原則として未登録ファイル一覧（Intake一覧）へ戻る。",
        ],
    )

    # 8. 注意事項
    add_heading_jp(doc, "8. 注意事項", 1)
    add_bullets(
        doc,
        [
            "添付なしで開始したフローでは、既存サービス案件へのファイル紐づけ（パターン7）はできません。ソースファイルが必須です。",
            "OCR読取には時間がかかることがあります。完了するまで画面を閉じず、結果を確認してから保存してください。",
            "削除ボタンは対象の未登録ファイルや関連データを削除します。取り消しできない場合があるため、実行前に対象を確認してください。",
            "Intake操作にはログイン（管理者権限）が必須です。セッション切れの場合は再ログインしてください。",
            "ファイル形式やサイズ制限がある場合は、システム設定に従ってください。",
        ],
    )

    # 9. 付録
    add_heading_jp(doc, "9. 付録：主要URL一覧", 1)
    add_para(doc, "ローカル開発環境を想定した主要URLです。環境によりパスは異なる場合があります。")
    add_table(
        doc,
        ["名称", "URL"],
        [
            ["ベース", BASE_URL],
            ["ホーム", f"{BASE_URL}/"],
            ["Intake（未登録ファイル一覧）", f"{BASE_URL}/（ホーム経由でIntakeへ）"],
        ],
    )
    add_para(
        doc,
        "実際のIntake画面URLはアプリケーションのルーティングに依存します。"
        "運用時はホーム画面のIntake導線から入る方法を推奨します。",
    )

    OUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUT_PATH))
    print(f"OUT={OUT_PATH}")
    print(f"SIZE={OUT_PATH.stat().st_size}")
    print("EMBEDDED:")
    for name in EMBEDDED:
        print(f"  {name}")
    print("MISSING:")
    if MISSING:
        for name in MISSING:
            print(f"  {name}")
    else:
        print("  (none)")
    print(f"JP_FONT={JP_FONT}")


if __name__ == "__main__":
    build()
